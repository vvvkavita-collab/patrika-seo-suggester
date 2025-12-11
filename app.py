# filename: patrika_seo_suggester_full.py
import streamlit as st
import requests
from bs4 import BeautifulSoup
import re
import json
import io
import csv
from datetime import datetime
from collections import Counter
from text_unidecode import unidecode
from docx import Document
import os
import html
import logging

# Optional: try to import newspaper (better extraction) if available
try:
    from newspaper import Article as NewspaperArticle  # pip install newspaper3k
    HAS_NEWSPAPER = True
except Exception:
    HAS_NEWSPAPER = False

# Optional: OpenAI (if user has key)
try:
    import openai
    HAS_OPENAI = True
except Exception:
    HAS_OPENAI = False

# -------------------------------
# Config / Defaults
# -------------------------------
st.set_page_config(page_title="SEO Sugester — Patrika", layout="wide")
PRIMARY_BRAND = "Rajasthan Patrika"
DEFAULT_AUTHOR = "Patrika News Desk"
DEFAULT_SECTION = "National"

HINDI_STOPWORDS = set("""
के की का हैं है और या यह था थी थे तथा लेकिन पर से में हो होना रहे रही रहे अगर तो भी लिए तक उन उस वही वहीँ एवं क्योंकि जैसे जैसेकि द्वारा नहीं बिना सभी उनका उनकी उनके वहीँ कभी हमेशा आदि प्रति लिए गए गई गया करें करेगा करेंगी करना करने करनेवाला करता करती करते जिसमें जिससे जिसके जिन जिसे जितना जितनी जितने आदि
""".split())
EN_STOPWORDS = set("""
the a an and or but if then else when while of for to in on at from by with without as is are was were be been being
""".split())

INTERNAL_LINKS = {
    "National": [
        ("Congress News", "https://www.patrika.com/national-news/congress/"),
        ("National Politics", "https://www.patrika.com/national-news/politics/"),
    ],
    "Rajasthan": [
        ("Jaipur News", "https://www.patrika.com/jaipur-news/"),
    ],
}

# -------------------------------
# Utilities (text processing + SEO heuristics)
# -------------------------------
def clean_text(txt: str) -> str:
    if not txt:
        return ""
    return " ".join(txt.replace("\r", "").split())

def tokenize(txt: str):
    tokens = []
    for w in txt.split():
        w = w.strip(".,:;!?'\"()[]{}“”‘’-–—|/\\")
        if w:
            tokens.append(w.lower())
    return tokens

def is_stopword(w: str) -> bool:
    return w in EN_STOPWORDS or w in HINDI_STOPWORDS

def top_keywords(text: str, n=6):
    tokens = tokenize(text)
    toks = [t for t in tokens if not is_stopword(t) and t.isalpha()]
    freq = Counter(toks)
    blacklist = {"news", "india", "indian", "said", "statement", "khabar", "patrika"}
    filtered = [(k, v) for k, v in freq.items() if k not in blacklist]
    filtered.sort(key=lambda x: (-x[1], x[0]))
    return [k for k, _ in filtered[:n]]

AUTHOR_NOISE = set(["by", "staff", "reporter", "updated", "written", "photo", "image", "author"])

def guess_primary_entity(text: str):
    # improved entity guess: find multi-word Titlecase sequences
    if not text:
        return "Breaking News"
    candidates = re.findall(r'\b([A-Z][a-zA-Z]+\s+[A-Z][a-zA-Z]+(?:\s+[A-Z][a-zA-Z]+)*)\b', text)
    singles = re.findall(r'\b([A-Z][a-zA-Z]{2,})\b', text)
    freq = Counter()
    for c in candidates:
        low = c.lower()
        if any(tok in low for tok in AUTHOR_NOISE):
            continue
        freq[c] += 2
    for s in singles:
        low = s.lower()
        if low in HINDI_STOPWORDS or low in EN_STOPWORDS or low in [a.lower() for a in AUTHOR_NOISE]:
            continue
        freq[s] += 1
    known = ["Shashi Tharoor", "Veer Savarkar", "Congress", "BJP", "Rajasthan", "Jaipur", "Delhi"]
    for k in known:
        if k.lower() in text.lower():
            return k
    if freq:
        return freq.most_common(1)[0][0]
    return "Breaking News"

def clamp(s: str, max_len: int):
    return s[:max_len].rstrip()

def clean_headline_text(title: str):
    if not title:
        return ""
    t = title.strip()
    # Remove common author suffix patterns and separators
    t = re.sub(r'\b[Bb]y[:\s]+[A-Z][\w\s\.-]{1,50}$', '', t).strip()
    t = re.split(r'[\|\—\–\:]{1,}.*$', t)[0].strip()
    t = re.sub(r'\b([Uu]pdated.*|[Ff]eatured.*|[Pp]hoto.*)$', '', t).strip()
    return t

def generate_title(body_text: str, original_title: str = None, max_len=60):
    if original_title:
        cleaned = clean_headline_text(original_title)
        if len(cleaned) >= 6:
            kws = top_keywords(body_text, n=2)
            if kws and kws[0] not in cleaned.lower():
                cand = f"{cleaned} — {kws[0].title()}"
                cleaned = cand if len(cand) <= max_len else cleaned
            return clamp(cleaned, max_len)
    sentences = re.split(r'(?<=[।\.\?\!])\s+', body_text.strip())
    first = sentences[0] if sentences else body_text.strip()
    words = first.split()
    short = " ".join(words[:12])
    kws = top_keywords(body_text, n=2)
    entity = guess_primary_entity(body_text)
    if entity != "Breaking News":
        title = f"{entity}: {short}"
    elif kws:
        title = f"{short} — {kws[0].title()}"
    else:
        title = short
    return clamp(title.strip(), max_len)

def generate_meta(body_text: str, max_len=160):
    body_clean = re.sub(r'\s+', ' ', body_text).strip()
    words = body_clean.split()
    take = 30 if len(words) >= 30 else len(words)
    snippet = " ".join(words[:take])
    kws = top_keywords(body_text, n=3)
    add = (" | Keywords: " + ", ".join(kws[:3])) if kws else ""
    meta = (snippet + add).strip()
    meta = re.sub(r'\b(By\s+[A-Z][\w\s\.]{1,40})$', '', meta).strip()
    return clamp(meta, max_len)

def slugify(title: str):
    ascii_title = unidecode(title)
    ascii_title = ascii_title.lower()
    allowed = []
    for ch in ascii_title:
        if ch.isalnum() or ch in [" ", "-", "_"]:
            allowed.append(ch)
    s = "".join(allowed).replace(" ", "-")
    parts = [p for p in s.split("-") if p and p not in EN_STOPWORDS]
    s = "-".join(parts)
    return clamp(s, 64)

def image_alts(text: str, count=2):
    kws = top_keywords(text, n=4)
    entity = guess_primary_entity(text)
    alts = []
    base = " ".join([entity] + kws[:2]).strip()
    if not base:
        base = "news image"
    for i in range(count):
        alts.append(clamp(f"{base} - scene {i+1}", 80))
    return alts

def readability_notes(text: str):
    notes = []
    paragraphs = [p.strip() for p in re.split(r'\n{1,}', text) if p.strip()]
    avg_len = sum(len(p) for p in paragraphs) / max(1, len(paragraphs))
    if avg_len > 500:
        notes.append("पैराग्राफ छोटे रखें (3–4 लाइन), लंबे पैराग्राफ विभाजित करें।")
    tokens = tokenize(text)
    if len(tokens) > 800:
        notes.append("इंट्रो छोटा करें और उपशीर्षक (H2/H3) जोड़कर सेक्शन्स बनाएं।")
    if not any(h in text for h in ["\n##", "\n###", "H2", "H3"]):
        notes.append("कम-से-कम 2 उपशीर्षक जोड़ें: पृष्ठभूमि, बयान/प्रतिक्रिया, संदर्भ।")
    if not notes:
        notes.append("रीडेबिलिटी ठीक है; छोटे पैराग्राफ और स्पष्ट उपशीर्षक बनाए रखें।")
    return notes

def schema_json_ld(headline, description, date_published, author, publisher, section, images=None):
    data = {
        "@context": "https://schema.org",
        "@type": "NewsArticle",
        "headline": headline,
        "description": description,
        "datePublished": date_published,
        "author": {"@type": "Person", "name": author},
        "publisher": {"@type": "Organization", "name": publisher},
        "articleSection": section,
        "isAccessibleForFree": True
    }
    if images:
        data["image"] = images
    return json.dumps(data, ensure_ascii=False, indent=2)

def html_snippet(title, meta, canonical, json_ld):
    return f"""<!-- SEO snippet start -->
<title>{html.escape(title)}</title>
<meta name="description" content="{html.escape(meta)}">
<link rel="canonical" href="{html.escape(canonical)}">
<script type="application/ld+json">
{json_ld}
</script>
<!-- SEO snippet end -->
"""

def docx_file(title, meta, body, keywords, slug, schema, links, alts, notes):
    doc = Document()
    doc.add_heading("SEO Suggester Output", level=1)
    doc.add_heading("Suggested Title", level=2)
    doc.add_paragraph(title)
    doc.add_heading("Suggested Meta", level=2)
    doc.add_paragraph(meta)
    doc.add_heading("Suggested Keywords", level=2)
    doc.add_paragraph(", ".join(keywords))
    doc.add_heading("Suggested URL Slug", level=2)
    doc.add_paragraph(slug)
    doc.add_heading("NewsArticle JSON-LD", level=2)
    doc.add_paragraph(schema)
    doc.add_heading("Suggested Internal Links", level=2)
    for text, url in links:
        doc.add_paragraph(f"{text}: {url}")
    doc.add_heading("Suggested Image Alt Text", level=2)
    for alt in alts:
        doc.add_paragraph(alt)
    doc.add_heading("Readability Notes", level=2)
    for n in notes:
        doc.add_paragraph(f"- {n}")
    doc.add_heading("Rewritten Article (Paragraph-wise)", level=2)
    for p in re.split(r'\n{1,}', body):
        if p.strip():
            doc.add_paragraph(p.strip())
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def csv_file_row(article_id, reporter, title, meta, slug, section):
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["ArticleID", "Reporter", "Title", "Meta", "Slug", "Section"])
    writer.writerow([article_id, reporter, title, meta, slug, section])
    output.seek(0)
    return output

# -------------------------------
# Fetching / Extraction
# -------------------------------
def fetch_article_from_url(url, timeout=12):
    """
    Try several heuristics to extract title/body. Return dict with keys:
    {title, body, canonical, error(optional)}
    """
    try:
        headers = {"User-Agent": "Mozilla/5.0 (compatible; SEO-Suggester/1.0)"}
        # Optionally use newspaper3k if available (better at extracting)
        if HAS_NEWSPAPER:
            try:
                art = NewspaperArticle(url)
                art.download()
                art.parse()
                title = art.title or ""
                body = art.text or ""
                canonical = art.source_url or url
                if body and len(body.split()) > 30:
                    return {"title": title or "", "body": body or "", "canonical": canonical}
            except Exception:
                # fallback to requests/BS
                pass

        resp = requests.get(url, headers=headers, timeout=timeout)
        resp.raise_for_status()
        soup = BeautifulSoup(resp.content, "html.parser")
        # title heuristics
        title = ""
        if soup.title and soup.title.string:
            title = soup.title.string.strip()
        h1 = soup.find("h1")
        if h1 and h1.get_text(strip=True):
            title = h1.get_text(strip=True)
        # article body heuristics
        article_text = ""
        article = soup.find("article")
        if article:
            ps = article.find_all("p")
            article_text = "\n\n".join([p.get_text(" ", strip=True) for p in ps if p.get_text(strip=True)])
        if not article_text:
            selectors = [
                {"name": "div", "attr": {"class": re.compile(r"(article|story|content|post|main|entry)", re.I)}},
                {"name": "div", "attr": {"id": re.compile(r"(article|story|content|main|entry)", re.I)}},
            ]
            for sel in selectors:
                nodes = soup.find_all(sel["name"], sel["attr"])
                for n in nodes:
                    ps = n.find_all("p")
                    if len(ps) >= 2:
                        article_text = "\n\n".join([p.get_text(" ", strip=True) for p in ps if p.get_text(strip=True)])
                        if len(article_text.split()) > 50:
                            break
                if article_text:
                    break
        if not article_text:
            ps = (soup.body.find_all("p") if soup.body else [])
            article_text = "\n\n".join([p.get_text(" ", strip=True) for p in ps if len(p.get_text(strip=True)) > 30])
        # canonical
        link_can = soup.find("link", {"rel": "canonical"})
        canonical = link_can["href"] if link_can and link_can.get("href") else url
        return {"title": title or "", "body": article_text or "", "canonical": canonical}
    except Exception as e:
        logging.exception("fetch error")
        return {"title": "", "body": "", "canonical": url, "error": str(e)}

# -------------------------------
# OpenAI integration (optional)
# -------------------------------
def get_openai_api_key():
    """
    Look for key in environment or streamlit secrets
    - Put in ~/.streamlit/secrets.toml as:
      [openai]
      api_key = "sk-..."
    Or set environment variable OPENAI_API_KEY
    """
    # streamlit secrets
    try:
        key = st.secrets["openai"]["api_key"]
        if key:
            return key
    except Exception:
        pass
    # env var
    key = os.environ.get("OPENAI_API_KEY") or os.environ.get("OPENAI_KEY")
    return key

def call_openai_rewrite(article_body: str, original_title: str = "", max_tokens=800):
    """
    Send a structured prompt to OpenAI to return JSON with:
    { title, meta, headings: [ {h2, h3s?}], paragraphs: [p1, p2, ...], keywords: [...], slug }
    Ensure model returns JSON ONLY.
    """
    api_key = get_openai_api_key()
    if not api_key or not HAS_OPENAI:
        return None  # caller will fallback to heuristics

    openai.api_key = api_key

    prompt = f"""
You are an expert news editor and SEO specialist for a Hindi+English newspaper. 
Input: raw news article body below, and optionally the scraped original title.
Task: produce a JSON object ONLY (no extra commentary) with these fields:
- title: a Google/SEO-friendly headline (Hindi or Hinglish allowed). 50-60 characters if possible.
- meta: an SEO meta description (150-160 chars ideal).
- slug: url-safe slug (lowercase, hyphen-separated).
- keywords: an array of 5 short keywords/phrases.
- headings: an array of section objects. Each section object: {{ "h2": "<H2 text>", "h3": ["sub1", "sub2", ...] }}. Provide at least 2 H2s if appropriate.
- paragraphs: an array of paragraph strings — rewrite the article into clear short paragraphs (3-6 lines each). Keep factual content same; do not invent facts. If facts missing, keep neutral.
- notes: short array of readability/SEO notes (2-4 items).

Constraints:
- Do NOT add new factual claims that aren't in the body. If something is unclear, keep it neutral (e.g., "प्रतिक्रिया का इंतजार है").
- Produce clean JSON only. Use Unicode (Hindi) where natural.
- Keep title and meta within recommended lengths (truncate if necessary).
- Use the original_title as hint for tone/subject if present.

Original title:
{original_title}

Article body:
{article_body}

Respond with JSON only.
"""

    try:
        response = openai.ChatCompletion.create(
            model="gpt-4o-mini",  # change as desired; ensure availability in your account
            messages=[{"role": "system", "content": "You are a helpful assistant."},
                      {"role": "user", "content": prompt}],
            max_tokens=max_tokens,
            temperature=0.2,
        )
        text = response["choices"][0]["message"]["content"].strip()
        # try to parse JSON even if model adds code fences
        json_text = re.sub(r"^```(?:json)?\s*", "", text)
        json_text = re.sub(r"\s*```$", "", json_text)
        parsed = json.loads(json_text)
        return parsed
    except Exception as e:
        logging.exception("OpenAI call failed")
        return None

# -------------------------------
# UI & Main
# -------------------------------
st.title("Patrika SEO Suggester — URL & Paste modes")
st.caption("Give published article URL(s) or paste article text. Outputs: SEO-friendly Title, Meta, Headings, Paragraph-wise rewrite. Optional OpenAI rewrite if API key available.")

with st.sidebar:
    st.header("Settings")
    section = st.selectbox("Article Section", ["National", "Rajasthan", "Business", "Sports", "Entertainment"], index=0)
    author = st.text_input("Author", value=DEFAULT_AUTHOR)
    publisher = st.text_input("Publisher", value=PRIMARY_BRAND)
    canonical_base = st.text_input("Canonical base URL (used if none found)", value="https://www.patrika.com")
    img_count = st.slider("Image alt suggestions (count)", 1, 4, 2)
    st.markdown("---")
    st.write("OpenAI / AI settings")
    st.write(f"OpenAI available in environment: {HAS_OPENAI and bool(get_openai_api_key())}")
    st.write("If you want better rewrites, set OPENAI_API_KEY in Streamlit secrets or environment variable.")
    st.markdown("---")
    st.write("Notes:")
    st.write("- For URL mode: paste one URL per line.")
    st.write("- For Paste mode: separate multiple articles with a blank line + --- + blank line (i.e. \\n\\n---\\n\\n).")

mode = st.radio("Select input mode", ["From URL(s)", "Paste Article(s)"])

def render_and_download_outputs(article_id, suggested, body_text, links):
    # suggested is dict with keys: title, meta, slug, keywords, headings, paragraphs, notes
    title = suggested.get("title") or ""
    meta = suggested.get("meta") or ""
    slug = suggested.get("slug") or slugify(title or "article")
    keywords = suggested.get("keywords") or top_keywords(body_text, n=6)
    headings = suggested.get("headings") or []
    paragraphs = suggested.get("paragraphs") or re.split(r'\n{1,}', body_text)
    notes = suggested.get("notes") or readability_notes(body_text)
    date_published = datetime.now().strftime("%Y-%m-%dT%H:%M:%S%z")
    schema = schema_json_ld(title, meta, date_published, author, publisher, section, images=None)
    canonical = canonical_base.rstrip("/") + "/" + section.lower() + "/" + slug

    st.markdown("### Suggested Title")
    st.write(title)
    st.markdown("### Suggested Meta")
    st.write(meta)
    st.markdown("### Suggested Keywords")
    st.write(", ".join(keywords))
    st.markdown("### Suggested URL Slug")
    st.code(slug, language="text")

    st.markdown("### Suggested Headings (H2/H3)")
    for sec in headings:
        h2 = sec.get("h2")
        h3s = sec.get("h3", [])
        st.write(f"- H2: {h2}")
        for h3 in h3s:
            st.write(f"  - H3: {h3}")

    st.markdown("### Rewritten Article (paragraph-wise)")
    for i, p in enumerate(paragraphs, start=1):
        st.write(f"**Paragraph {i}:** {p}")

    st.markdown("### Readability / SEO Notes")
    for n in notes:
        st.write(f"- {n}")

    st.markdown("### NewsArticle JSON-LD")
    st.code(schema, language="json")

    snippet = html_snippet(title, meta, canonical, schema)
    st.markdown("### HTML snippet")
    st.code(snippet, language="html")

    # Downloads
    docx_bytes = docx_file(title, meta, "\n\n".join(paragraphs), keywords, slug, schema, links, image_alts(body_text, img_count), notes)
    st.download_button("Download DOCX", data=docx_bytes, file_name=f"{article_id}_seo_suggestions.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    st.download_button("Download JSON-LD", data=schema, file_name=f"{article_id}_newsarticle.json", mime="application/ld+json")
    st.download_button("Download HTML snippet", data=snippet, file_name=f"{article_id}_seo_snippet.html", mime="text/html")
    csv_io = csv_file_row(article_id, author, title, meta, slug, section)
    st.download_button("Download CSV (management)", data=csv_io.getvalue(), file_name=f"{article_id}_summary.csv", mime="text/csv")

if mode == "From URL(s)":
    st.subheader("Paste published article URL(s) — one URL per line")
    urls_input = st.text_area("URLs (one per line)", height=160, placeholder="https://www.example.com/news/123\nhttps://another.example.com/story/abc")
    if st.button("Fetch & Analyze URLs", disabled=len(urls_input.strip()) == 0):
        urls = [u.strip() for u in urls_input.splitlines() if u.strip()]
        if not urls:
            st.warning("Please paste at least one URL.")
        else:
            for idx, url in enumerate(urls, start=1):
                st.markdown(f"---\n## Article {idx}: `{url}`")
                with st.spinner(f"Fetching {url} ..."):
                    fetched = fetch_article_from_url(url)
                # DEBUG: show raw extraction to catch author-in-title issues
                st.markdown("**[Debug] Raw extraction preview**")
                st.write("Extracted title (raw):", fetched.get("title", "")[:300])
                st.write("Extracted body (raw preview, first 700 chars):")
                st.code((fetched.get("body", "")[:700] + ("..." if len(fetched.get("body",""))>700 else "")))
                if fetched.get("error"):
                    st.error(f"Error fetching {url}: {fetched['error']}")
                    continue
                body = clean_text(fetched.get("body", ""))
                fetched_title = fetched.get("title", "").strip()
                canonical_url = fetched.get("canonical") or canonical_base.rstrip("/") + "/" + section.lower()
                if not body or len(body.split()) < 30:
                    st.warning("Unable to extract a full article body from this URL. You can paste the article manually in Paste mode.")
                    if fetched_title:
                        st.write("Extracted title (partial):")
                        st.write(fetched_title)
                    continue

                # try OpenAI rewrite first (if available)
                suggested = None
                if HAS_OPENAI and get_openai_api_key():
                    with st.spinner("Calling OpenAI to rewrite (if available)..."):
                        suggested = call_openai_rewrite(body, original_title=fetched_title)
                if not suggested:
                    # fallback heuristics
                    title = generate_title(body, original_title=fetched_title)
                    meta = generate_meta(body)
                    slug = slugify(title)
                    keywords = top_keywords(body, n=6)
                    headings = [
                        {"h2": "पृष्ठभूमि", "h3": []},
                        {"h2": "बयान / प्रतिक्रिया", "h3": []}
                    ]
                    paragraphs = [p.strip() for p in re.split(r'\n{1,}', body) if p.strip()]
                    suggested = {
                        "title": title,
                        "meta": meta,
                        "slug": slug,
                        "keywords": keywords,
                        "headings": headings,
                        "paragraphs": paragraphs,
                        "notes": readability_notes(body)
                    }

                article_id = f"URLART-{datetime.now().strftime('%Y%m%d%H%M%S')}-{idx}"
                links = INTERNAL_LINKS.get(section, [])
                render_and_download_outputs(article_id, suggested, body, links)

elif mode == "Paste Article(s)":
    st.subheader("Paste your news article(s)")
    st.info("If multiple articles: separate them with a blank line + --- + blank line (i.e. \\n\\n---\\n\\n).")
    news_text = st.text_area("Paste full body (headline optional).", height=300, placeholder="अपनी खबर यहाँ पेस्ट करें...\n\n---\n\n(Next article)")
    if st.button("Analyze & Suggest (Paste mode)", disabled=len(news_text.strip()) == 0):
        parts = [p.strip() for p in re.split(r'\n{0,}\-{3,}\n{0,}', news_text) if p.strip()]
        if not parts:
            st.warning("No article found—please paste correctly.")
        else:
            for idx, part in enumerate(parts, start=1):
                st.markdown(f"---\n## Pasted Article {idx}")
                body = clean_text(part)
                # If user pasted "headline\n\nbody" we can try to split
                possible_title = ""
                lines = [l.strip() for l in part.splitlines() if l.strip()]
                if len(lines) >= 2 and len(lines[0].split()) <= 12:
                    # treat first line as possible original title
                    possible_title = lines[0]

                suggested = None
                if HAS_OPENAI and get_openai_api_key():
                    with st.spinner("Calling OpenAI to rewrite (if available)..."):
                        suggested = call_openai_rewrite(body, original_title=possible_title)
                if not suggested:
                    title = generate_title(body, original_title=possible_title)
                    meta = generate_meta(body)
                    slug = slugify(title)
                    keywords = top_keywords(body, n=6)
                    headings = [
                        {"h2": "पृष्ठभूमि", "h3": []},
                        {"h2": "बयान / प्रतिक्रिया", "h3": []}
                    ]
                    paragraphs = [p.strip() for p in re.split(r'\n{1,}', body) if p.strip()]
                    suggested = {
                        "title": title,
                        "meta": meta,
                        "slug": slug,
                        "keywords": keywords,
                        "headings": headings,
                        "paragraphs": paragraphs,
                        "notes": readability_notes(body)
                    }
                article_id = f"PASTEART-{datetime.now().strftime('%Y%m%d%H%M%S')}-{idx}"
                links = INTERNAL_LINKS.get(section, [])
                render_and_download_outputs(article_id, suggested, body, links)

# Footer
st.markdown("---")
st.caption("Outputs are heuristic and editor-friendly. Editor should review before publishing. If URL extraction fails, use Paste mode. For more robust extraction on JS-heavy sites, consider enabling newspaper3k or a headless browser.")
